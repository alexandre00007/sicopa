import tempfile
import unittest
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader, PdfWriter
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas

from controle_paie.file_tools import pdf_to_excel, pdf_to_word, rotate_pdf


class FileToolsTests(unittest.TestCase):
    def setUp(self):
        self.temp = tempfile.TemporaryDirectory()
        self.root = Path(self.temp.name)
        self.source = self.root / "source.pdf"
        pdf = canvas.Canvas(str(self.source))
        pdf.drawString(72, 760, "Matricule Nom Montant")
        pdf.drawString(72, 740, "A001 Alexandre 1500")
        pdf.showPage()
        pdf.drawString(72, 760, "Deuxieme page")
        pdf.save()

    def tearDown(self):
        self.temp.cleanup()

    def test_rotate_pdf_preserves_pages_and_reports_progress(self):
        target = self.root / "rotation.pdf"; events = []
        rotate_pdf(str(self.source), str(target), 90, lambda value, label: events.append((value, label)))
        reader = PdfReader(str(target))
        self.assertEqual(len(reader.pages), 2)
        self.assertEqual(reader.pages[0].rotation, 90)
        self.assertEqual(events[-1][0], 100)

    def test_pdf_to_excel_extracts_text_with_summary(self):
        target = self.root / "conversion.xlsx"
        pdf_to_excel(str(self.source), str(target))
        workbook = load_workbook(target, read_only=True)
        self.assertEqual(workbook.sheetnames[0], "Sommaire")
        values = [cell for row in workbook["Page_1"].iter_rows(values_only=True) for cell in row if cell]
        self.assertTrue(any("Alexandre" in str(value) for value in values))
        workbook.close()

    def test_pdf_to_word_extracts_text(self):
        target = self.root / "conversion.docx"
        pdf_to_word(str(self.source), str(target))
        document = Document(target)
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("Alexandre", text)
        self.assertIn("Page 2", text)

    def test_scanned_pdf_is_automatically_recognized(self):
        image = Image.new("RGB", (1800, 500), "white")
        draw = ImageDraw.Draw(image)
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 72)
        draw.text((80, 90), "MATRICULE A001", fill="black", font=font)
        draw.text((80, 210), "ALEXANDRE MULUMBA", fill="black", font=font)
        scanned = self.root / "document_scane.pdf"
        pdf = canvas.Canvas(str(scanned), pagesize=(900, 250))
        pdf.drawImage(ImageReader(image), 0, 0, width=900, height=250)
        pdf.save()
        target = self.root / "ocr.xlsx"; events = []
        pdf_to_excel(str(scanned), str(target), lambda value, label: events.append((value, label)))
        workbook = load_workbook(target, read_only=True)
        values = " ".join(str(cell) for row in workbook["Page_1"].iter_rows(values_only=True) for cell in row if cell)
        summary = " ".join(str(cell) for row in workbook["Sommaire"].iter_rows(values_only=True) for cell in row if cell)
        workbook.close()
        self.assertIn("MATRICULE", values)
        self.assertIn("MULUMBA", values)
        self.assertIn("Texte OCR", summary)
        self.assertTrue(any("OCR automatique" in label for _, label in events))
        word_target = self.root / "ocr.docx"
        pdf_to_word(str(scanned), str(word_target))
        word_text = "\n".join(paragraph.text for paragraph in Document(word_target).paragraphs)
        self.assertIn("MULUMBA", word_text)
        self.assertIn("obtenu par OCR", word_text)

    def test_image_only_pdf_requests_ocr(self):
        source = self.root / "scan.pdf"; writer = PdfWriter(); writer.add_blank_page(300, 300)
        with source.open("wb") as stream: writer.write(stream)
        with self.assertRaisesRegex(ValueError, "OCR"):
            pdf_to_excel(str(source), str(self.root / "scan.xlsx"))


if __name__ == "__main__":
    unittest.main()
