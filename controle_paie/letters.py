from __future__ import annotations

import os
from datetime import datetime
from decimal import Decimal
from pathlib import Path


def _money(value) -> str:
    amount=Decimal(str(value or 0))
    return f"{amount:,.2f}".replace(","," ").replace(".",",")


def _integer(value) -> str:
    return f"{int(value or 0):,}".replace(","," ")


def _find(summaries: list[dict], *parts: str) -> dict:
    for item in summaries:
        label=item["label"].lower()
        if all(part.lower() in label for part in parts):return item
    return {"count":0,"concerned":0,"impact":0,"label":""}


def generate_interpretation_letter(path: Path, institution: str, regime: str, quarter: str,
                                   year: int, summaries: list[dict], filters: list[tuple]) -> Path:
    """Create a formal Word interpretation letter next to the final report."""
    try:
        from docx import Document
        from docx.enum.section import WD_SECTION
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:
        raise RuntimeError("La génération de la lettre Word nécessite python-docx. Installez les dépendances avec : python -m pip install -r requirements.txt") from exc

    NAVY=RGBColor(18,53,91);BLUE=RGBColor(22,119,255);INK=RGBColor(36,50,71);MUTED=RGBColor(97,113,135);WHITE=RGBColor(255,255,255)
    doc=Document();section=doc.sections[0]
    section.page_width=Inches(8.5);section.page_height=Inches(11)
    section.top_margin=Inches(1);section.bottom_margin=Inches(1);section.left_margin=Inches(1);section.right_margin=Inches(1)
    section.header_distance=Inches(.492);section.footer_distance=Inches(.492)
    doc.core_properties.title=f"Lettre d’interprétation — {institution} — {quarter} {year}"
    doc.core_properties.subject="Interprétation du contrôle et du rapprochement de la paie"
    doc.core_properties.author="SICORPA — Alexandre Mulumba Kande"
    doc.core_properties.keywords="SICORPA, contrôle de la paie, rapprochement, interprétation"

    def font(run,size=11,bold=False,color=INK,italic=False):
        run.font.name="Calibri";run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"),"Calibri");run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"),"Calibri")
        run.font.size=Pt(size);run.bold=bold;run.italic=italic;run.font.color.rgb=color;return run
    normal=doc.styles["Normal"];normal.font.name="Calibri";normal.font.size=Pt(11);normal.font.color.rgb=INK
    normal.paragraph_format.space_before=Pt(0);normal.paragraph_format.space_after=Pt(6);normal.paragraph_format.line_spacing=1.10
    for name,size,before,after in [("Heading 1",16,12,6),("Heading 2",13,12,6),("Heading 3",12,8,4)]:
        style=doc.styles[name];style.font.name="Calibri";style.font.size=Pt(size);style.font.bold=True;style.font.color.rgb=BLUE
        style.paragraph_format.space_before=Pt(before);style.paragraph_format.space_after=Pt(after);style.paragraph_format.keep_with_next=True

    header=section.header.paragraphs[0];header.alignment=WD_ALIGN_PARAGRAPH.LEFT
    font(header.add_run("SICORPA"),9.5,True,NAVY);font(header.add_run("  |  Lettre d’interprétation du contrôle de la paie"),9,False,MUTED)
    footer=section.footer.paragraphs[0];footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
    font(footer.add_run("SICORPA • Document généré automatiquement • "),8,False,MUTED)
    fld=OxmlElement("w:fldSimple");fld.set(qn("w:instr"),"PAGE");footer._p.append(fld)

    p=doc.add_paragraph();p.paragraph_format.space_after=Pt(2);font(p.add_run("SICORPA"),23,True,NAVY)
    p=doc.add_paragraph();p.paragraph_format.space_after=Pt(16);font(p.add_run("Système Intégré de Contrôle et de Rapprochement de la Paie"),12,False,MUTED)
    for label,value in [("À",institution),("Période",f"{quarter} {year}"),("Régime",regime),("Date",datetime.now().strftime("%d/%m/%Y")),("Objet","Transmission et interprétation des résultats du contrôle de la paie")]:
        p=doc.add_paragraph();p.paragraph_format.space_after=Pt(2);font(p.add_run(f"{label} : "),10.5,True,NAVY);font(p.add_run(str(value)),10.5)
    p=doc.add_paragraph();p.paragraph_format.space_before=Pt(12);p.paragraph_format.space_after=Pt(8);font(p.add_run("Madame, Monsieur le Responsable,"),11,True,NAVY)
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    font(p.add_run(f"Dans le cadre du contrôle de la paie de {institution}, SICORPA a analysé le régime {regime} pour la période {quarter} {year}. La présente lettre accompagne le rapport final et ses annexes. Elle synthétise les principaux constats techniques sans se substituer à la validation administrative, financière ou disciplinaire de l’institution."))

    listing=_find(summaries,"données du listing")
    declaratif=_find(summaries,"données déclaratives")
    missing=_find(summaries,"matricules manquants")
    dup_mat=_find(summaries,"doublons","matricule")
    dup_name=_find(summaries,"doublons","nom")
    present=_find(summaries,"déclarés présents","matricule")
    unpaid_declared=_find(summaries,"non déclarés","matricule")
    elsewhere=_find(summaries,"payés hors périmètre","matricule")

    doc.add_heading("1. Périmètre analysé",level=1)
    table=doc.add_table(rows=1,cols=3);table.alignment=WD_TABLE_ALIGNMENT.LEFT;table.autofit=False
    widths=[3120,1800,4440];headers=["Élément","Volume","Lecture"]
    for i,(cell,width,title) in enumerate(zip(table.rows[0].cells,widths,headers)):
        cell.width=Inches(width/1440);cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade=OxmlElement("w:shd");shade.set(qn("w:fill"),"12355B");cell._tc.get_or_add_tcPr().append(shade);font(cell.paragraphs[0].add_run(title),10,True,WHITE)
    perimeter=[("Listing de paie filtré",listing.get("count",0),"Lignes soumises au contrôle"),("Liste déclarative",declaratif.get("count",0),"Lignes transmises par l’institution"),("Filtres métier",len(filters or []),"Règles appliquées avant comparaison")]
    for label,value,reading in perimeter:
        cells=table.add_row().cells
        for cell,width in zip(cells,widths):cell.width=Inches(width/1440);cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        font(cells[0].paragraphs[0].add_run(label),10,True,NAVY);cells[1].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER;font(cells[1].paragraphs[0].add_run(_integer(value)),10,True,NAVY);font(cells[2].paragraphs[0].add_run(reading),10)
    _set_table_geometry(table,widths,qn,OxmlElement)

    doc.add_heading("2. Principaux constats",level=1)
    findings=[
        ("Matricules manquants ou non exploitables",missing,"Ces lignes réduisent la fiabilité d’un rapprochement strict par identifiant."),
        ("Doublons par matricule",dup_mat,"Ils signalent plusieurs lignes rattachées à un même matricule valide; NU et N.U sont exclus."),
        ("Doublons par nom",dup_name,"Ils constituent un indice complémentaire et ne doivent pas être additionnés aux doublons par matricule."),
        ("Déclarés présents dans le listing filtré",present,"Cette population forme le noyau confirmé par rapprochement au matricule."),
        ("Non déclarés mais présents dans le listing",unpaid_declared,"Ces cas nécessitent la vérification des pièces individuelles et du périmètre déclaratif."),
        ("Agents payés hors du périmètre filtré",elsewhere,"Ces agents déclarés et présents dans le listing filtré ont aussi été retrouvés ailleurs dans le listing trimestriel."),
    ]
    ftable=doc.add_table(rows=1,cols=4);ftable.alignment=WD_TABLE_ALIGNMENT.LEFT;ftable.autofit=False
    fwidths=[3440,1200,1600,3120]
    for cell,width,title in zip(ftable.rows[0].cells,fwidths,["Rubrique","Concernés","Impact potentiel","Interprétation"]):
        cell.width=Inches(width/1440);cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade=OxmlElement("w:shd");shade.set(qn("w:fill"),"E8EEF5");cell._tc.get_or_add_tcPr().append(shade);font(cell.paragraphs[0].add_run(title),9.3,True,NAVY)
    for label,item,meaning in findings:
        cells=ftable.add_row().cells
        for cell,width in zip(cells,fwidths):cell.width=Inches(width/1440);cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        font(cells[0].paragraphs[0].add_run(label),9.2,True,NAVY);cells[1].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER;font(cells[1].paragraphs[0].add_run(_integer(item.get("concerned",0))),9.2,True,NAVY)
        cells[2].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT;font(cells[2].paragraphs[0].add_run(_money(item.get("impact",0))),9.2)
        font(cells[3].paragraphs[0].add_run(meaning),9.0)
    _set_table_geometry(ftable,fwidths,qn,OxmlElement)

    doc.add_heading("3. Interprétation générale",level=1)
    anomaly_total=sum(int(item.get("concerned",0)) for _,item,_ in findings[:3])
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if anomaly_total:
        font(p.add_run("Le contrôle met en évidence des cas nécessitant une revue ciblée. "),11,True,NAVY)
        font(p.add_run("Les volumes signalés sont des indicateurs techniques. Une même personne peut apparaître dans les contrôles par matricule et par nom; ces effectifs ne doivent donc pas être additionnés pour établir un total global de personnes en anomalie."))
    else:
        font(p.add_run("Aucune anomalie majeure n’a été détectée dans les rubriques prioritaires analysées. "),11,True,NAVY)
        font(p.add_run("Cette conclusion reste limitée au périmètre, aux filtres et à la qualité des données chargées."))
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    font(p.add_run("Les montants indiqués comme impacts potentiels représentent une exposition technique calculée à partir des rémunérations associées aux lignes contrôlées. Ils ne constituent ni une créance certaine ni un manque à gagner définitivement établi avant confirmation documentaire et décision de l’autorité compétente."))
    formulas=[]
    for item in summaries:
        formula=item.get("formula",{})
        if formula and formula.get("id") not in {entry[0] for entry in formulas}:formulas.append((formula.get("id"),formula.get("name","Formule"),formula.get("version",1)))
    formula_trace=(" Formules appliquées : "+"; ".join(f"{name} (v{version})" for _identifier,name,version in formulas)+"; détails dans la feuille « Formules d’impact ».") if formulas else ""

    doc.add_heading("4. Recommandations",level=1)
    recommendations=[
        "Examiner prioritairement les fichiers d’effectifs uniques liés aux rubriques comportant des concernés.",
        "Faire confirmer les matricules, affectations et pièces justificatives par les services compétents.",
        "Analyser séparément les résultats par matricule et par nom afin d’éviter tout double comptage.",
        "Documenter la validation ou le rejet de chaque anomalie avant de confirmer un impact financier.",
        "Conserver le rapport final, les annexes et la présente lettre dans le dossier de contrôle de la période.",
    ]
    for text in recommendations:
        p=doc.add_paragraph(style="List Bullet");p.paragraph_format.left_indent=Inches(.5);p.paragraph_format.first_line_indent=Inches(-.25);font_space = Pt(5);p.paragraph_format.space_after=font_space;p.paragraph_format.line_spacing=1.167;font(p.add_run(text))

    doc.add_heading("5. Documents joints",level=1)
    p=doc.add_paragraph();font(p.add_run("Le dossier comprend le rapport_final.xlsx, les annexes détaillées et les fichiers d’effectifs uniques. Les liens intégrés au rapport final permettent d’ouvrir les pièces correspondantes."+formula_trace))
    p=doc.add_paragraph();p.paragraph_format.space_before=Pt(10);p.paragraph_format.space_after=Pt(4);p.paragraph_format.keep_with_next=True;font(p.add_run("Veuillez agréer, Madame, Monsieur le Responsable, l’expression de notre considération distinguée."))
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.RIGHT;p.paragraph_format.space_before=Pt(8);p.paragraph_format.space_after=Pt(0);p.paragraph_format.keep_together=True
    font(p.add_run("SICORPA"),11,True,NAVY);p.add_run().add_break();font(p.add_run("Système de contrôle et de rapprochement de la paie"),9.5,False,MUTED)

    temporary=path.with_name(f".{path.stem}.part.docx")
    try:doc.save(temporary);os.replace(temporary,path)
    finally:
        if temporary.exists():temporary.unlink()
    return path


def _set_table_geometry(table, widths: list[int], qn, OxmlElement) -> None:
    """Apply fixed 9360-DXA geometry and comfortable cell margins."""
    table_xml=table._tbl;properties=table_xml.tblPr
    width=properties.find(qn("w:tblW"))
    if width is None:width=OxmlElement("w:tblW");properties.append(width)
    width.set(qn("w:type"),"dxa");width.set(qn("w:w"),str(sum(widths)))
    indent=properties.find(qn("w:tblInd"))
    if indent is None:indent=OxmlElement("w:tblInd");properties.append(indent)
    indent.set(qn("w:type"),"dxa");indent.set(qn("w:w"),"120")
    grid=table_xml.tblGrid
    for child in list(grid):grid.remove(child)
    for value in widths:
        column=OxmlElement("w:gridCol");column.set(qn("w:w"),str(value));grid.append(column)
    for row_index,row in enumerate(table.rows):
        row_props=row._tr.get_or_add_trPr()
        cannot_split=OxmlElement("w:cantSplit");row_props.append(cannot_split)
        if row_index==0:
            repeat=OxmlElement("w:tblHeader");repeat.set(qn("w:val"),"true");row_props.append(repeat)
        for cell,value in zip(row.cells,widths):
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before=0;paragraph.paragraph_format.space_after=0;paragraph.paragraph_format.line_spacing=1.05
            props=cell._tc.get_or_add_tcPr();tcw=props.find(qn("w:tcW"))
            if tcw is None:tcw=OxmlElement("w:tcW");props.append(tcw)
            tcw.set(qn("w:type"),"dxa");tcw.set(qn("w:w"),str(value))
            margins=props.find(qn("w:tcMar"))
            if margins is None:margins=OxmlElement("w:tcMar");props.append(margins)
            for edge,amount in [("top",80),("bottom",80),("start",120),("end",120)]:
                node=margins.find(qn(f"w:{edge}"))
                if node is None:node=OxmlElement(f"w:{edge}");margins.append(node)
                node.set(qn("w:w"),str(amount));node.set(qn("w:type"),"dxa")
