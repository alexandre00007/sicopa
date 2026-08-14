from __future__ import annotations

import json
import re
import uuid
from datetime import datetime
from pathlib import Path
from typing import Callable, Optional

from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill

from .database import Database
from .spreadsheet_utils import sanitize_excel_row, sanitize_xml_text

Progress = Optional[Callable[[int, str], None]]


class ListingGroupAnalysisService:
    """Constitue et analyse une base issue uniquement de listings de paie."""

    RUBRICS = {
        "MATRICULE_NON_EXPLOITABLE": "MATRICULE_MANQUANT",
        "PAYE_PLUSIEURS_REGIMES": "PAYE_HORS_PERIMETRE",
        "PAYE_PLUSIEURS_INSTITUTIONS": "PAYE_HORS_PERIMETRE",
        "DOUBLON_MATRICULE": "DOUBLON_MATRICULE",
        "DOUBLON_NOM": "DOUBLON_NOM",
    }

    def __init__(self,database: Database):
        self.db=database

    def available_sources(self,quarter: str,year: int) -> list[tuple]:
        with self.db.connect() as con:
            return con.execute("""SELECT p.execution_id,COALESCE(i.nom_officiel,p.institution_id),
                    p.institution_id,p.regime,p.table_source,COUNT(*),MIN(j.fichier_source)
                FROM paie_standardisee p
                LEFT JOIN institutions i ON i.institution_id=p.institution_id
                LEFT JOIN journal_executions j ON j.execution_id=p.execution_id
                WHERE p.trimestre=? AND p.annee=?
                GROUP BY p.execution_id,i.nom_officiel,p.institution_id,p.regime,p.table_source
                ORDER BY p.regime,i.nom_officiel,p.table_source""",[quarter,int(year)]).fetchall()

    def preview(self,quarter: str,year: int,execution_ids: list[str]) -> list[dict]:
        available={row[0]:row for row in self.available_sources(quarter,year)}
        result=[]
        with self.db.connect() as con:
            for execution_id in execution_ids:
                source=available.get(execution_id)
                if not source:continue
                _,name,institution_id,regime,table_source,count,file_source=source
                clause,params=self.db.payroll_filter_clause(institution_id,regime,"p")
                retained=con.execute(f"""SELECT COUNT(*) FROM paie_standardisee p
                    WHERE p.execution_id=?{clause}""",[execution_id]+params).fetchone()[0]
                filters=self.db.list_treatment_filters(institution_id,regime)
                quality=con.execute("""SELECT
                    SUM(CASE WHEN matricule_normalise NOT IN ('','NU') THEN 1 ELSE 0 END),
                    SUM(CASE WHEN nom_normalise<>'' THEN 1 ELSE 0 END)
                    FROM paie_standardisee WHERE execution_id=?""",[execution_id]).fetchone()
                issues=[]
                if retained==0:issues.append("Aucune ligne après filtres")
                if int(quality[0] or 0)+int(quality[1] or 0)==0:issues.append("Aucun identifiant exploitable")
                result.append({"execution_id":execution_id,"institution":name,
                    "institution_id":institution_id,"regime":regime,"table":table_source,
                    "file":file_source,"available":count,"retained":retained,
                    "filters":filters,"ready":not issues,"issues":issues})
        return result

    def sample_source(self,execution_id: str,limit: int=50) -> tuple[list[str],list[tuple]]:
        with self.db.connect() as con:
            source=con.execute("""SELECT institution_id,regime FROM paie_standardisee
                WHERE execution_id=? LIMIT 1""",[execution_id]).fetchone()
            if not source:raise ValueError("Source introuvable.")
            clause,params=self.db.payroll_filter_clause(source[0],source[1],"p")
            columns=["matricule_source","nom","prenom","section","categorie","grade",
                     "unite_affectation","province","remuneration_brute_calculee","montant_net"]
            rows=con.execute(f"""SELECT {','.join('p.'+c for c in columns)}
                FROM paie_standardisee p WHERE p.execution_id=?{clause}
                ORDER BY p.ligne_source LIMIT ?""",
                [execution_id]+params+[max(1,min(int(limit),200))]).fetchall()
        return columns,rows

    def run(self,name: str,quarter: str,year: int,execution_ids: list[str],
            progress: Progress=None) -> dict:
        if not name.strip():raise ValueError("Donnez un nom au groupe de listings.")
        if not execution_ids:raise ValueError("Sélectionnez au moins une source.")
        available={row[0]:row for row in self.available_sources(quarter,year)}
        if any(item not in available for item in execution_ids):
            raise ValueError("Une source ne correspond plus à la période choisie.")
        diagnostics=self.preview(quarter,year,execution_ids)
        if not diagnostics or any(not item["ready"] for item in diagnostics):
            raise ValueError("Corrigez les sources signalées avant de constituer le groupe.")
        group_id=str(uuid.uuid4());progress and progress(5,"Création du groupe analytique")
        with self.db.connect() as con:
            con.execute("""INSERT INTO groupes_analyse_listing
                VALUES (?,?,?,?,'EN_COURS',0,CURRENT_TIMESTAMP,NULL,FALSE,NULL)""",
                [group_id,name.strip(),quarter,int(year)])
            con.execute("BEGIN")
            try:
                total=max(1,len(execution_ids))
                for index,execution_id in enumerate(execution_ids,1):
                    source=available[execution_id]
                    _,_label,institution_id,regime,table_source,count,_file=source
                    filters=self.db.list_treatment_filters(institution_id,regime)
                    clause,params=self.db.payroll_filter_clause(institution_id,regime,"p")
                    con.execute(f"""INSERT INTO base_analyse_listing BY NAME
                        SELECT ? groupe_id,p.* FROM paie_standardisee p
                        WHERE p.execution_id=? AND p.trimestre=? AND p.annee=?{clause}""",
                        [group_id,execution_id,quarter,int(year)]+params)
                    retained=con.execute("""SELECT COUNT(*) FROM base_analyse_listing
                        WHERE groupe_id=? AND execution_id=?""",[group_id,execution_id]).fetchone()[0]
                    filter_text=json.dumps([{"colonne":f[1],"operateur":f[2],"valeur":f[3]}
                                            for f in filters],ensure_ascii=False)
                    con.execute("INSERT INTO sources_analyse_listing VALUES (?,?,?,?,?,?,?,?)",
                        [group_id,execution_id,institution_id,regime,table_source,count,retained,filter_text])
                    progress and progress(10+int(25*index/total),
                        f"Source {index}/{total} : {regime} — {retained:,} lignes".replace(","," "))

                progress and progress(40,"Calcul des occurrences et des croisements")
                con.execute("""INSERT INTO resultats_analyse_listing
                    WITH keyed AS (
                        SELECT b.*,CASE
                          WHEN matricule_normalise NOT IN ('','NU') THEN 'M:'||matricule_normalise
                          WHEN nom_normalise<>'' THEN 'N:'||nom_normalise
                          ELSE 'L:'||ligne_paie_id END person_key
                        FROM base_analyse_listing b WHERE groupe_id=?
                    ),stats AS (
                        SELECT *,
                          CASE WHEN matricule_normalise NOT IN ('','NU')
                            THEN COUNT(*) OVER(PARTITION BY matricule_normalise) ELSE 0 END occ_mat,
                          CASE WHEN nom_normalise<>''
                            THEN COUNT(*) OVER(PARTITION BY nom_normalise) ELSE 0 END occ_nom,
                          COUNT(DISTINCT regime) OVER(PARTITION BY person_key) nb_regimes,
                          COUNT(DISTINCT institution_id) OVER(PARTITION BY person_key) nb_institutions,
                          ROW_NUMBER() OVER(PARTITION BY person_key
                            ORDER BY regime,institution_id,ligne_source) rang
                        FROM keyed
                    )
                    SELECT uuid(),?,ligne_paie_id,institution_id,regime,execution_id,table_source,
                      CASE
                        WHEN matricule_normalise IN ('','NU') THEN 'MATRICULE_NON_EXPLOITABLE'
                        WHEN nb_regimes>1 THEN 'PAYE_PLUSIEURS_REGIMES'
                        WHEN nb_institutions>1 THEN 'PAYE_PLUSIEURS_INSTITUTIONS'
                        WHEN occ_mat>1 THEN 'DOUBLON_MATRICULE'
                        WHEN occ_nom>1 THEN 'DOUBLON_NOM'
                        ELSE 'UNIQUE_DANS_GROUPE' END,
                      occ_mat,occ_nom,nb_regimes,nb_institutions,rang,
                      remuneration_brute_calculee,0,'AUCUN_IMPACT',
                      matricule_normalise,nom_normalise
                    FROM stats""",[group_id,group_id])
                progress and progress(60,"Calcul des impacts selon chaque régime source")
                for index,execution_id in enumerate(execution_ids,1):
                    source=available[execution_id]
                    _,_label,institution_id,regime,_table,_count,_file=source
                    expressions={};formula_ids={}
                    for status,rubric in self.RUBRICS.items():
                        expression,formula=self.db.impact_sql(
                            institution_id,regime,quarter,int(year),rubric,"b","r.rang_occurrence")
                        expressions[status]=expression
                        formula_ids[status]=str(formula["id"]).replace("'","''")
                    impact_case="CASE r.statut_analyse "+" ".join(
                        f"WHEN '{status}' THEN CAST(({expressions[status]}) AS DECIMAL(38,2))"
                        for status in self.RUBRICS)+" ELSE 0 END"
                    formula_case="CASE r.statut_analyse "+" ".join(
                        f"WHEN '{status}' THEN '{formula_ids[status]}'"
                        for status in self.RUBRICS)+" ELSE 'AUCUN_IMPACT' END"
                    con.execute(f"""UPDATE resultats_analyse_listing r SET
                        impact_potentiel={impact_case},formule_impact_id={formula_case}
                        FROM base_analyse_listing b
                        WHERE r.groupe_id=? AND r.ligne_paie_id=b.ligne_paie_id
                          AND b.execution_id=?""",[group_id,execution_id])
                    progress and progress(60+int(25*index/total),
                        f"Impacts {index}/{total} : régime {regime}")
                base_count=con.execute("SELECT COUNT(*) FROM base_analyse_listing WHERE groupe_id=?",
                                       [group_id]).fetchone()[0]
                con.execute("""UPDATE groupes_analyse_listing SET statut='TERMINE',
                    lignes_base=?,termine_le=CURRENT_TIMESTAMP WHERE groupe_id=?""",
                    [base_count,group_id])
                con.execute("COMMIT")
            except Exception:
                con.execute("ROLLBACK")
                con.execute("""UPDATE groupes_analyse_listing SET statut='ERREUR',
                    termine_le=CURRENT_TIMESTAMP WHERE groupe_id=?""",[group_id])
                raise
        progress and progress(100,"Analyse groupée des listings terminée")
        return {"group_id":group_id,"base_rows":base_count,"summary":self.summary(group_id)}

    def summary(self,group_id: str) -> list[tuple]:
        with self.db.connect() as con:
            return con.execute("""SELECT statut_analyse,COUNT(*),
                    COUNT(DISTINCT COALESCE(NULLIF(matricule_normalise,''),
                        NULLIF(nom_normalise,''),ligne_paie_id)),
                    COALESCE(SUM(masse_financiere),0),COALESCE(SUM(impact_potentiel),0)
                FROM resultats_analyse_listing WHERE groupe_id=?
                GROUP BY statut_analyse ORDER BY statut_analyse""",[group_id]).fetchall()

    def inter_regime_summary(self,group_id: str) -> list[tuple]:
        with self.db.connect() as con:
            return con.execute("""SELECT regime,statut_analyse,COUNT(*),
                    COUNT(DISTINCT COALESCE(NULLIF(matricule_normalise,''),
                        NULLIF(nom_normalise,''),ligne_paie_id)),
                    COALESCE(SUM(masse_financiere),0),COALESCE(SUM(impact_potentiel),0)
                FROM resultats_analyse_listing WHERE groupe_id=?
                GROUP BY regime,statut_analyse ORDER BY regime,statut_analyse""",[group_id]).fetchall()

    def list_groups(self,include_archived: bool=False) -> list[tuple]:
        where="" if include_archived else "WHERE NOT COALESCE(archive,FALSE)"
        with self.db.connect() as con:
            return con.execute(f"""SELECT groupe_id,nom,trimestre,annee,statut,lignes_base,
                    cree_le,termine_le,dossier_export,COALESCE(archive,FALSE)
                FROM groupes_analyse_listing {where} ORDER BY cree_le DESC""").fetchall()

    def archive_group(self,group_id: str,archived: bool=True) -> None:
        with self.db.connect() as con:
            if not con.execute("SELECT 1 FROM groupes_analyse_listing WHERE groupe_id=?",
                               [group_id]).fetchone():raise ValueError("Groupe introuvable.")
            con.execute("UPDATE groupes_analyse_listing SET archive=? WHERE groupe_id=?",
                        [archived,group_id])

    def export(self,group_id: str,root: str,progress: Progress=None) -> Path:
        with self.db.connect() as con:
            group=con.execute("""SELECT nom,trimestre,annee,lignes_base FROM groupes_analyse_listing
                WHERE groupe_id=?""",[group_id]).fetchone()
        if not group:raise ValueError("Groupe introuvable.")
        name,quarter,year,base_rows=group
        slug=re.sub(r"[^A-Za-z0-9]+","_",name).strip("_").lower()
        folder=Path(root)/f"analyse_listings_{year}_{quarter}_{slug}_{datetime.now():%Y%m%d_%H%M%S}"
        folder.mkdir(parents=True,exist_ok=True);progress and progress(5,"Création du rapport")

        report=Workbook();ws=report.active;ws.title="Synthèse"
        ws.append(["ANALYSE GROUPÉE DES LISTINGS"]);ws["A1"].font=Font(bold=True,size=16,color="12355B")
        ws.append(sanitize_excel_row(["Groupe",name]));ws.append(sanitize_excel_row(["Période",f"{quarter} {year}"]));ws.append(["Lignes",base_rows]);ws.append([])
        ws.append(["Catégorie","Enregistrements","Concernés","Masse financière","Impact potentiel"])
        for cell in ws[6]:cell.font=Font(bold=True,color="FFFFFF");cell.fill=PatternFill("solid",fgColor="12355B")
        for row in self.summary(group_id):ws.append(sanitize_excel_row(row))
        ws.freeze_panes="A7";ws.auto_filter.ref=f"A6:E{ws.max_row}"
        for col,width in zip("ABCDE",[43,18,18,22,22]):ws.column_dimensions[col].width=width
        methods=report.create_sheet("Méthodologie")
        for line in [
            "Le groupe contient uniquement les listings sélectionnés de la même période.",
            "Les filtres métier sont appliqués séparément à chaque institution et régime.",
            "NU, N.U et leurs variantes sont classés comme matricules non exploitables.",
            "La classification recherche ensuite les paiements multi-régimes, multi-institutions, les doublons de matricule puis de nom.",
            "Chaque ligne reçoit une seule classification finale.",
            "L'impact est calculé avec la formule du régime et de l'institution de la ligne source.",
        ]:methods.append([line])
        report_path=folder/"rapport_analyse_listings.xlsx"

        headers=["Statut","Institution","Régime","Table source","Matricule","Nom complet","Nom normalisé",
                 "Occurrences matricule","Occurrences nom","Régimes","Institutions","Rang",
                 "Masse financière","Impact potentiel","Formule"]
        detail_book=Workbook(write_only=True);detail=detail_book.create_sheet("Résultats")
        detail.append(headers)
        query="""SELECT r.statut_analyse,i.nom_officiel,r.regime,r.table_source,
                r.matricule_normalise,
                TRIM(CONCAT_WS(' ', COALESCE(NULLIF(b.nom,''),''), COALESCE(NULLIF(b.prenom,''),''))) AS nom_complet,
                r.nom_normalise,r.occurrences_matricule,
                r.occurrences_nom,r.nombre_regimes,r.nombre_institutions,r.rang_occurrence,
                r.masse_financiere,r.impact_potentiel,r.formule_impact_id
            FROM resultats_analyse_listing r
            LEFT JOIN institutions i ON i.institution_id=r.institution_id
            LEFT JOIN base_analyse_listing b
              ON b.groupe_id=r.groupe_id AND b.ligne_paie_id=r.ligne_paie_id
            WHERE r.groupe_id=? ORDER BY r.statut_analyse,r.regime,r.nom_normalise"""
        with self.db.connect() as con:
            cursor=con.execute(query,[group_id]);written=0
            while True:
                rows=cursor.fetchmany(5000)
                if not rows:break
                for row in rows:detail.append(sanitize_excel_row(row))
                written+=len(rows);progress and progress(min(55,20+int(35*written/max(1,base_rows))),
                    f"Annexe globale : {written:,} lignes".replace(","," "))
        detail_path=folder/"annexe_globale_listings.xlsx";detail_book.save(detail_path)
        progress and progress(57,f"Fichier généré : {detail_path.name}")

        category_folder=folder/"annexes_par_categorie";category_folder.mkdir(exist_ok=True)
        categories=[row[0] for row in self.summary(group_id)]
        for index,status in enumerate(categories,1):
            book=Workbook(write_only=True);sheet=book.create_sheet("Détails");sheet.append(headers)
            with self.db.connect() as con:
                cursor=con.execute(query.replace("WHERE r.groupe_id=?","WHERE r.groupe_id=? AND r.statut_analyse=?"),
                                   [group_id,status])
                while True:
                    rows=cursor.fetchmany(5000)
                    if not rows:break
                    for row in rows:sheet.append(sanitize_excel_row(row))
            filename=status.lower()+".xlsx";book.save(category_folder/filename)
            progress and progress(57+int(18*index/max(1,len(categories))),
                f"Fichier généré : annexes_par_categorie/{filename}")

        people=Workbook(write_only=True);sheet=people.create_sheet("Effectifs uniques")
        sheet.append(["Clé agent","Matricule","Nom complet","Nom normalisé","Catégorie","Régimes","Institutions",
                      "Occurrences","Masse","Impact"])
        with self.db.connect() as con:
            cursor=con.execute("""SELECT CASE WHEN r.matricule_normalise NOT IN ('','NU')
                            THEN 'M:'||r.matricule_normalise ELSE 'N:'||r.nom_normalise END cle,
                    MAX(r.matricule_normalise),
                    MAX(TRIM(CONCAT_WS(' ', COALESCE(NULLIF(b.nom,''),''), COALESCE(NULLIF(b.prenom,''),'')))),
                    MAX(r.nom_normalise),MAX(r.statut_analyse),
                    STRING_AGG(DISTINCT r.regime,', '),STRING_AGG(DISTINCT i.nom_officiel,', '),
                    COUNT(*),SUM(r.masse_financiere),SUM(r.impact_potentiel)
                FROM resultats_analyse_listing r
                LEFT JOIN institutions i ON i.institution_id=r.institution_id
                LEFT JOIN base_analyse_listing b
                  ON b.groupe_id=r.groupe_id AND b.ligne_paie_id=r.ligne_paie_id
                WHERE r.groupe_id=? GROUP BY cle ORDER BY cle""",[group_id])
            while True:
                rows=cursor.fetchmany(5000)
                if not rows:break
                for row in rows:sheet.append(sanitize_excel_row(row))
        people_path=folder/"effectifs_uniques_listings.xlsx";people.save(people_path)
        progress and progress(85,f"Fichier généré : {people_path.name}")

        letter=Document();letter.styles["Normal"].font.name="Arial";letter.styles["Normal"].font.size=Pt(10)
        title=letter.add_paragraph();run=title.add_run("LETTRE D’INTERPRÉTATION — ANALYSE GROUPÉE DES LISTINGS")
        run.bold=True;run.font.size=Pt(15)
        letter.add_paragraph(sanitize_xml_text(f"Groupe : {name} — Période : {quarter} {year}"))
        letter.add_paragraph(f"La base constituée contient {base_rows:,} lignes après application des filtres propres à chaque source.".replace(","," "))
        table=letter.add_table(rows=1,cols=4);table.style="Table Grid"
        for cell,value in zip(table.rows[0].cells,["Catégorie","Concernés","Masse","Impact"]):cell.text=value
        for status,_records,concerned,mass,impact in self.summary(group_id):
            cells=table.add_row().cells
            for cell,value in zip(cells,[status,concerned,f"{mass:,.2f}".replace(","," "),
                                         f"{impact:,.2f}".replace(","," ")]):cell.text=sanitize_xml_text(value)
        letter.add_paragraph("Les impacts sont potentiels et doivent être confirmés par les pièces administratives et la validation métier.")
        letter_path=folder/"lettre_interpretation_listings.docx";letter.save(letter_path)
        progress and progress(95,f"Fichier généré : {letter_path.name}")
        ws.append([]);ws.append(["Liens annexes","Annexe globale","Effectifs uniques","Lettre d’interprétation"])
        links_row=ws.max_row
        for col,target in [(2,"annexe_globale_listings.xlsx"),(3,"effectifs_uniques_listings.xlsx"),(4,"lettre_interpretation_listings.docx")]:
            cell=ws.cell(links_row,col);cell.value=target;cell.hyperlink=target;cell.style="Hyperlink"
        for cell in ws[links_row]:
            cell.font=Font(bold=True,color="12355B")
        report.save(report_path)
        progress and progress(20,f"Fichier généré : {report_path.name}")
        with self.db.connect() as con:
            con.execute("UPDATE groupes_analyse_listing SET dossier_export=? WHERE groupe_id=?",
                        [str(folder),group_id])
        progress and progress(100,"Rapport, annexes, effectifs et lettre terminés")
        return folder
