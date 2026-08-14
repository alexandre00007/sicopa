from __future__ import annotations

import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path
from typing import Callable, Optional

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from pypdf import PdfReader, PdfWriter

from .spreadsheet_utils import sanitize_excel_row, sanitize_xml_text


Progress = Optional[Callable[[int, str], None]]


def tesseract_path() -> Optional[Path]:
    """Locate an installed or application-bundled Tesseract executable."""
    candidates = []
    bundled = getattr(sys, "_MEIPASS", None)
    if bundled:
        candidates.extend([Path(bundled) / "tesseract" / "tesseract.exe", Path(bundled) / "tesseract"])
    detected = shutil.which("tesseract")
    if detected:
        candidates.append(Path(detected))
    if sys.platform == "win32":
        candidates.extend([Path(r"C:\Program Files\Tesseract-OCR\tesseract.exe"), Path(r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe")])
    return next((candidate for candidate in candidates if candidate.is_file()), None)


def ocr_available() -> bool:
    return tesseract_path() is not None


def _ocr_page(page, language: str = "fra+eng") -> str:
    executable = tesseract_path()
    if not executable:
        raise ValueError("Ce PDF nécessite un OCR, mais Tesseract n’est pas installé. Installez Tesseract OCR avec la langue française puis relancez la conversion.")
    try:
        from PIL import ImageOps
        image = page.to_image(resolution=250, antialias=True).original
        image = ImageOps.autocontrast(image.convert("L"))
        with tempfile.TemporaryDirectory(prefix="sicorpa_ocr_") as folder:
            image_path = Path(folder) / "page.png"; image.save(image_path, "PNG")
            result = subprocess.run([str(executable), str(image_path), "stdout", "-l", language, "--psm", "6"], capture_output=True, text=True, timeout=240, check=False)
    except subprocess.TimeoutExpired as exc:
        raise ValueError("L’OCR a dépassé quatre minutes sur une page. Réduisez la résolution du PDF puis réessayez.") from exc
    except Exception as exc:
        raise ValueError(f"Impossible de préparer la page pour l’OCR : {exc}") from exc
    if result.returncode:
        detail=(result.stderr or "Erreur Tesseract inconnue").strip()
        raise ValueError(f"Échec de l’OCR ({language}) : {detail}")
    return result.stdout.strip()


def _paths(source: str, target: str, suffix: str) -> tuple[Path, Path]:
    source_path = Path(source).expanduser().resolve()
    target_path = Path(target).expanduser().resolve()
    if not source_path.is_file() or source_path.suffix.lower() != ".pdf":
        raise ValueError("Sélectionnez un fichier PDF valide.")
    if target_path.suffix.lower() != suffix:
        raise ValueError(f"Le fichier de destination doit porter l’extension {suffix}.")
    if source_path == target_path:
        raise ValueError("Le fichier source et le fichier de destination doivent être différents.")
    target_path.parent.mkdir(parents=True, exist_ok=True)
    return source_path, target_path


def _reader(source: Path) -> PdfReader:
    reader = PdfReader(str(source))
    if reader.is_encrypted:
        try:
            unlocked = reader.decrypt("")
        except Exception as exc:
            raise ValueError("Ce PDF est protégé par un mot de passe.") from exc
        if not unlocked:
            raise ValueError("Ce PDF est protégé par un mot de passe.")
    return reader


def rotate_pdf(source: str, target: str, degrees: int = 90, progress: Progress = None) -> Path:
    source_path, target_path = _paths(source, target, ".pdf")
    if degrees not in {90, 180, 270}:
        raise ValueError("La rotation doit être de 90°, 180° ou 270°.")
    reader = _reader(source_path)
    if not reader.pages:
        raise ValueError("Le PDF ne contient aucune page.")
    writer = PdfWriter(); total = len(reader.pages)
    for index, page in enumerate(reader.pages, 1):
        writer.add_page(page.rotate(degrees))
        progress and progress(int(index / total * 90), f"Rotation : page {index}/{total}")
    temporary = target_path.with_name(f".{target_path.name}.tmp")
    try:
        with temporary.open("wb") as stream:
            writer.write(stream)
        temporary.replace(target_path)
    finally:
        if temporary.exists():
            temporary.unlink()
    progress and progress(100, f"Fichier généré : {target_path.name}")
    return target_path


def _safe_sheet_name(value: str, used: set[str]) -> str:
    base = re.sub(r"[\\/*?:\[\]]", "_", value).strip()[:31] or "Feuille"
    candidate = base; number = 2
    while candidate in used:
        suffix = f"_{number}"; candidate = f"{base[:31-len(suffix)]}{suffix}"; number += 1
    used.add(candidate)
    return candidate


def _style_sheet(sheet) -> None:
    sheet.freeze_panes = "A2"; sheet.sheet_view.showGridLines = False
    for cell in sheet[1]:
        cell.font = Font(name="Aptos", size=10, bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="12355B")
        cell.alignment = Alignment(vertical="center", wrap_text=True)
    sheet.row_dimensions[1].height = 25
    for column in range(1, sheet.max_column + 1):
        values = [str(sheet.cell(row, column).value or "") for row in range(1, min(sheet.max_row, 250) + 1)]
        sheet.column_dimensions[get_column_letter(column)].width = min(55, max(12, max(map(len, values), default=10) + 2))
        for row in range(2, sheet.max_row + 1):
            sheet.cell(row, column).alignment = Alignment(vertical="top", wrap_text=True)
    if sheet.max_row > 1:
        sheet.auto_filter.ref = sheet.dimensions


def pdf_to_excel(source: str, target: str, progress: Progress = None, use_ocr: bool = True, ocr_language: str = "fra+eng") -> Path:
    import pdfplumber

    source_path, target_path = _paths(source, target, ".xlsx")
    workbook = Workbook(); workbook.remove(workbook.active); used: set[str] = set(); index_rows = []
    extracted_content = 0; temporary = target_path.with_name(f".{target_path.stem}.tmp{target_path.suffix}")
    try:
        with pdfplumber.open(str(source_path)) as pdf:
            if not pdf.pages:
                raise ValueError("Le PDF ne contient aucune page.")
            total = len(pdf.pages)
            for page_number, page in enumerate(pdf.pages, 1):
                tables = page.extract_tables() or []
                valid_tables = [table for table in tables if table and any(any(cell not in (None, "") for cell in row) for row in table)]
                if valid_tables:
                    for table_number, table in enumerate(valid_tables, 1):
                        title = _safe_sheet_name(f"P{page_number}_Tableau{table_number}", used)
                        sheet = workbook.create_sheet(title)
                        for row in table:
                            sheet.append(sanitize_excel_row(cell if cell is not None else "" for cell in row))
                        _style_sheet(sheet); extracted_content += 1
                        index_rows.append([page_number, "Tableau", title, max(0, len(table) - 1)])
                else:
                    text = (page.extract_text() or "").strip(); ocr_used = False
                    if not text and use_ocr:
                        progress and progress(max(1,int((page_number-1)/total*90)),f"OCR automatique : page {page_number}/{total}")
                        text = _ocr_page(page,ocr_language); ocr_used = True
                    if text:
                        title = _safe_sheet_name(f"Page_{page_number}", used); sheet = workbook.create_sheet(title)
                        sheet.append(["Ligne", "Texte extrait"])
                        for line_number, line in enumerate(text.splitlines(), 1):
                            sheet.append(sanitize_excel_row([line_number, line]))
                        _style_sheet(sheet); extracted_content += 1
                        index_rows.append([page_number, "Texte OCR — à vérifier" if ocr_used else "Texte", title, len(text.splitlines())])
                progress and progress(int(page_number / total * 90), f"Extraction Excel : page {page_number}/{total}")
        if not extracted_content:
            raise ValueError("Aucun texte n’a été reconnu, même après l’OCR. Vérifiez la qualité du scan et la langue sélectionnée." if use_ocr else "Aucun texte n’a été détecté. Activez l’OCR pour traiter ce PDF scanné.")
        summary = workbook.create_sheet("Sommaire", 0)
        summary.append(["Page PDF", "Contenu détecté", "Feuille Excel", "Nombre de lignes"])
        for row in index_rows: summary.append(row)
        _style_sheet(summary)
        workbook.save(temporary); temporary.replace(target_path)
    except Exception:
        if temporary.exists(): temporary.unlink()
        raise
    finally:
        workbook.close()
    progress and progress(100, f"Fichier généré : {target_path.name}")
    return target_path


def _configure_document(document: Document, source_name: str) -> None:
    section = document.sections[0]
    section.top_margin = section.bottom_margin = Inches(0.8)
    section.left_margin = section.right_margin = Inches(0.85)
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"; normal.font.size = Pt(10)
    title = document.add_paragraph()
    run = title.add_run("Conversion du document PDF")
    run.bold = True; run.font.size = Pt(20); run.font.color.rgb = RGBColor(18, 53, 91)
    subtitle = document.add_paragraph(sanitize_xml_text(f"Source : {source_name}"))
    subtitle.runs[0].font.color.rgb = RGBColor(97, 113, 135)


def pdf_to_word(source: str, target: str, progress: Progress = None, use_ocr: bool = True, ocr_language: str = "fra+eng") -> Path:
    import pdfplumber

    source_path, target_path = _paths(source, target, ".docx")
    document = Document(); _configure_document(document, source_path.name); extracted_content = 0
    temporary = target_path.with_name(f".{target_path.stem}.tmp{target_path.suffix}")
    try:
        with pdfplumber.open(str(source_path)) as pdf:
            if not pdf.pages:
                raise ValueError("Le PDF ne contient aucune page.")
            total = len(pdf.pages)
            for page_number, page in enumerate(pdf.pages, 1):
                document.add_heading(f"Page {page_number}", level=1)
                text = (page.extract_text() or "").strip(); ocr_used = False
                if not text and use_ocr:
                    progress and progress(max(1,int((page_number-1)/total*90)),f"OCR automatique : page {page_number}/{total}")
                    text = _ocr_page(page,ocr_language); ocr_used = True
                if text:
                    if ocr_used:
                        warning=document.add_paragraph("Texte obtenu par OCR — vérifiez les matricules, noms et montants.")
                        warning.runs[0].italic=True;warning.runs[0].font.color.rgb=RGBColor(180,83,9)
                    for block in re.split(r"\n\s*\n", text):
                        document.add_paragraph(sanitize_xml_text(block.replace("\n", " ")))
                    extracted_content += 1
                tables = page.extract_tables() or []
                for table_data in tables:
                    if not table_data: continue
                    width = max((len(row) for row in table_data), default=0)
                    if not width: continue
                    table = document.add_table(rows=0, cols=width); table.style = "Table Grid"
                    for row_index, row_data in enumerate(table_data):
                        cells = table.add_row().cells
                        for column, value in enumerate(row_data): cells[column].text = sanitize_xml_text(value or "")
                        if row_index == 0:
                            for cell in cells:
                                for run in cell.paragraphs[0].runs: run.bold = True
                    extracted_content += 1
                if page_number < total:
                    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
                progress and progress(int(page_number / total * 90), f"Extraction Word : page {page_number}/{total}")
        if not extracted_content:
            raise ValueError("Aucun texte n’a été reconnu, même après l’OCR. Vérifiez la qualité du scan et la langue sélectionnée." if use_ocr else "Aucun texte n’a été détecté. Activez l’OCR pour traiter ce PDF scanné.")
        document.save(temporary); temporary.replace(target_path)
    except Exception:
        if temporary.exists(): temporary.unlink()
        raise
    progress and progress(100, f"Fichier généré : {target_path.name}")
    return target_path
