from __future__ import annotations

import json
import re
import uuid
from datetime import datetime
from pathlib import Path
from typing import Callable, Optional

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill
from docx import Document
from docx.shared import Pt

from .database import Database
from .spreadsheet_utils import sanitize_excel_row, sanitize_xml_text

Progress = Optional[Callable[[int, str], None]]


class MultiRegimeAnalysisService:
    """Analyse un déclaratif contre un instantané filtré de plusieurs listings."""

    def __init__(self, database: Database):
        self.db = database

    def available_declarations(self,institution_id: str,regime: str,quarter: str,year: int) -> list[tuple]:
        with self.db.connect() as con:
            return con.execute("""SELECT d.execution_id,MIN(d.fichier_source),MIN(d.feuille_source),
                    COUNT(*),MIN(j.date_fin)
                FROM declaratif_standardise d LEFT JOIN journal_executions j
                  ON j.execution_id=d.execution_id
                WHERE d.institution_id=? AND d.regime=? AND d.trimestre=? AND d.annee=?
                GROUP BY d.execution_id ORDER BY MIN(j.date_fin) DESC NULLS LAST,d.execution_id""",
                [institution_id,regime,quarter,int(year)]).fetchall()

    def _resolve_declaration(self,institution_id: str,regime: str,quarter: str,year: int,
                             execution_id: str="") -> tuple[str,int]:
        rows=self.available_declarations(institution_id,regime,quarter,year)
        if execution_id:
            selected=[row for row in rows if row[0]==execution_id]
            if not selected:raise ValueError("La version déclarative sélectionnée n'appartient plus à ce périmètre.")
            return execution_id,int(selected[0][3])
        if not rows:raise ValueError("Aucun déclaratif n'est chargé pour ce périmètre.")
        if len(rows)>1:raise ValueError("Plusieurs versions du déclaratif existent. Sélectionnez précisément celle à analyser.")
        return rows[0][0],int(rows[0][3])

    def diagnose(self,institution_id: str,declaration_regime: str,quarter: str,year: int,
                 declaration_execution_id: str,execution_ids: list[str]) -> dict:
        declaration_execution_id,declaration_count=self._resolve_declaration(
            institution_id,declaration_regime,quarter,year,declaration_execution_id)
        preview_rows=self.preview(quarter,year,execution_ids)
        sources=[]
        for execution_id,name,source_institution,source_regime,table_source,available,retained in preview_rows:
            filters=self.db.list_treatment_filters(source_institution,source_regime)
            formula=self.db.resolve_impact_formula(source_institution,source_regime,quarter,year,"PAYE_HORS_PERIMETRE")
            with self.db.connect() as con:
                quality=con.execute("""SELECT COUNT(*),
                        SUM(CASE WHEN matricule_normalise NOT IN ('','NU') THEN 1 ELSE 0 END),
                        SUM(CASE WHEN nom_normalise<>'' THEN 1 ELSE 0 END)
                    FROM paie_standardisee WHERE execution_id=?""",[execution_id]).fetchone()
                operation=con.execute("""SELECT type_operation FROM journal_executions
                    WHERE execution_id=? ORDER BY date_fin DESC LIMIT 1""",[execution_id]).fetchone()
            source_type={"IMPORT_ACCESS":"ACCESS","IMPORT_PAIE_EXCEL":"PAIE_EXCEL"}.get(
                operation[0] if operation else "","STANDARD")
            mapping_count=len(self.db.list_column_mappings(source_regime,source_type)) if source_type!="STANDARD" else 0
            mapping_label=(f"{source_type} — {mapping_count} mapping(s)"
                           if mapping_count else f"{source_type} — auto-détection validée")
            usable=int(quality[1] or 0)+int(quality[2] or 0)
            issues=[]
            if retained==0:issues.append("Aucune ligne après filtres")
            if usable==0:issues.append("Ni matricule ni nom exploitable")
            sources.append({"execution_id":execution_id,"institution":name,"institution_id":source_institution,
                "regime":source_regime,"table":table_source,"available":available,"retained":retained,
                "filters":filters,"formula":formula,"mapping":mapping_label,
                "ready":not issues,"issues":issues})
        return {"declaration_execution_id":declaration_execution_id,
                "declaration_rows":declaration_count,"sources":sources,
                "available_rows":sum(x["available"] for x in sources),
                "retained_rows":sum(x["retained"] for x in sources),
                "ready":bool(sources) and all(x["ready"] for x in sources)}

    def available_sources(self, quarter: str, year: int) -> list[tuple]:
        with self.db.connect() as con:
            return con.execute(
                """SELECT p.execution_id,COALESCE(i.nom_officiel,p.institution_id),p.institution_id,
                          p.regime,p.table_source,COUNT(*),MIN(j.fichier_source)
                   FROM paie_standardisee p
                   LEFT JOIN institutions i ON i.institution_id=p.institution_id
                   LEFT JOIN journal_executions j ON j.execution_id=p.execution_id
                   WHERE p.trimestre=? AND p.annee=?
                   GROUP BY p.execution_id,i.nom_officiel,p.institution_id,p.regime,p.table_source
                   ORDER BY p.regime,i.nom_officiel,p.table_source""",
                [quarter,int(year)]).fetchall()

    def preview(self, quarter: str, year: int, execution_ids: list[str]) -> list[tuple]:
        available={row[0]:row for row in self.available_sources(quarter,year)}
        rows=[]
        with self.db.connect() as con:
            for execution_id in execution_ids:
                source=available.get(execution_id)
                if not source:continue
                _,name,institution_id,regime,table_source,count,_file=source
                clause,params=self.db.payroll_filter_clause(institution_id,regime,"p")
                retained=con.execute(
                    f"SELECT COUNT(*) FROM paie_standardisee p WHERE p.execution_id=?{clause}",
                    [execution_id]+params).fetchone()[0]
                rows.append((execution_id,name,institution_id,regime,table_source,count,retained))
        return rows

    def sample_source(self,execution_id: str,limit: int=50) -> tuple[list[str],list[tuple]]:
        with self.db.connect() as con:
            source=con.execute("""SELECT institution_id,regime FROM paie_standardisee
                WHERE execution_id=? LIMIT 1""",[execution_id]).fetchone()
            if not source:raise ValueError("Source de paie introuvable.")
            institution_id,regime=source
            clause,params=self.db.payroll_filter_clause(institution_id,regime,"p")
            columns=["matricule_source","nom","prenom","section","categorie","grade",
                     "unite_affectation","province","remuneration_brute_calculee","montant_net"]
            rows=con.execute(f"""SELECT {','.join('p.'+column for column in columns)}
                FROM paie_standardisee p WHERE p.execution_id=?{clause}
                ORDER BY p.ligne_source LIMIT ?""",
                [execution_id]+params+[max(1,min(int(limit),200))]).fetchall()
        return columns,rows

    def run(self, institution_id: str, declaration_regime: str, quarter: str, year: int,
            execution_ids: list[str], declaration_execution_id: str="", progress: Progress=None) -> dict:
        if not execution_ids:raise ValueError("Sélectionnez au moins une source de paie.")
        available={row[0]:row for row in self.available_sources(quarter,year)}
        if any(item not in available for item in execution_ids):
            raise ValueError("Une source sélectionnée n'appartient plus à la période choisie.")
        declaration_execution_id,declaration_count=self._resolve_declaration(
            institution_id,declaration_regime,quarter,year,declaration_execution_id)

        campaign_id=str(uuid.uuid4())
        progress and progress(5,"Création de la campagne d'analyse")
        with self.db.connect() as con:
            con.execute("""INSERT INTO campagnes_analyse_multi
                (campagne_id,institution_declarative_id,regime_declaratif,trimestre,annee,
                 statut,lignes_base,lignes_declaratives,cree_le,termine_le,
                 declaratif_execution_id,archivee,dossier_export)
                VALUES (?,?,?,?,?,'EN_COURS',0,?,CURRENT_TIMESTAMP,NULL,?,FALSE,NULL)""",
                [campaign_id,institution_id,declaration_regime,quarter,int(year),
                 declaration_count,declaration_execution_id])
            con.execute("BEGIN")
            try:
                total=max(1,len(execution_ids))
                for index,execution_id in enumerate(execution_ids,1):
                    source=available[execution_id]
                    _,_name,source_institution,source_regime,table_source,count,_file=source
                    filters=self.db.list_treatment_filters(source_institution,source_regime)
                    clause,params=self.db.payroll_filter_clause(source_institution,source_regime,"p")
                    con.execute(f"""INSERT INTO base_analyse_multi BY NAME
                        SELECT ? campagne_id,p.* FROM paie_standardisee p
                        WHERE p.execution_id=? AND p.trimestre=? AND p.annee=?{clause}""",
                        [campaign_id,execution_id,quarter,int(year)]+params)
                    retained=con.execute("""SELECT COUNT(*) FROM base_analyse_multi
                        WHERE campagne_id=? AND execution_id=?""",
                        [campaign_id,execution_id]).fetchone()[0]
                    filter_text=json.dumps(
                        [{"colonne":f[1],"operateur":f[2],"valeur":f[3]} for f in filters],
                        ensure_ascii=False)
                    con.execute("INSERT INTO sources_analyse_multi VALUES (?,?,?,?,?,?,?,?)",
                        [campaign_id,execution_id,source_institution,source_regime,
                         table_source,count,retained,filter_text])
                    progress and progress(10+int(30*index/total),
                        f"Source {index}/{total} : {source_regime} — {retained:,} lignes retenues".replace(","," "))

                progress and progress(45,"Rapprochement avec la base trimestrielle")
                for index,execution_id in enumerate(execution_ids,1):
                    source=available[execution_id]
                    _,_name,source_institution,source_regime,_table,_count,_file=source
                    outside=source_institution!=institution_id or source_regime!=declaration_regime
                    if outside:
                        expression,formula=self.db.impact_sql(
                            source_institution,source_regime,quarter,int(year),
                            "PAYE_HORS_PERIMETRE","m","m.occurrence_rank")
                        formula_id=formula["id"]
                    else:
                        expression,formula_id="0","AUCUN_IMPACT"
                    con.execute(f"""INSERT INTO resultats_analyse_multi
                        WITH ds AS (
                            SELECT ligne_declaratif_id,matricule_normalise,nom_normalise
                            FROM declaratif_standardise
                            WHERE institution_id=? AND regime=? AND trimestre=? AND annee=?
                              AND execution_id=?
                        ),
                        matricule_hits AS (
                            SELECT DISTINCT d.ligne_declaratif_id
                            FROM ds d JOIN base_analyse_multi b
                              ON b.campagne_id=? AND d.matricule_normalise NOT IN ('','NU')
                             AND b.matricule_normalise=d.matricule_normalise
                        ),
                        matches_raw AS (
                            SELECT d.ligne_declaratif_id declaration_id,
                                   d.matricule_normalise declaration_matricule,
                                   d.nom_normalise declaration_nom,b.*,'MATRICULE' methode
                            FROM ds d JOIN base_analyse_multi b
                              ON b.campagne_id=? AND b.execution_id=?
                             AND d.matricule_normalise NOT IN ('','NU')
                             AND b.matricule_normalise=d.matricule_normalise
                            UNION ALL
                            SELECT d.ligne_declaratif_id,d.matricule_normalise,d.nom_normalise,
                                   b.*,'NOM' methode
                            FROM ds d LEFT JOIN matricule_hits h
                              ON h.ligne_declaratif_id=d.ligne_declaratif_id
                            JOIN base_analyse_multi b
                              ON b.campagne_id=? AND b.execution_id=?
                             AND d.nom_normalise<>'' AND b.nom_normalise=d.nom_normalise
                            WHERE h.ligne_declaratif_id IS NULL
                        ),
                        matched AS (
                            SELECT *,ROW_NUMBER() OVER(PARTITION BY declaration_id
                                ORDER BY regime,institution_id,ligne_source) occurrence_rank
                            FROM matches_raw
                        )
                        SELECT uuid(),?,m.declaration_id,m.ligne_paie_id,?,?,m.institution_id,
                               m.regime,m.execution_id,m.table_source,m.methode,'EN_COURS',0,0,
                               m.remuneration_brute_calculee,
                               CAST(({expression}) AS DECIMAL(38,2)),?,
                               m.declaration_matricule,m.declaration_nom
                        FROM matched m""",
                        [institution_id,declaration_regime,quarter,int(year),declaration_execution_id,
                         campaign_id,campaign_id,execution_id,campaign_id,execution_id,
                         campaign_id,institution_id,declaration_regime,formula_id])
                    progress and progress(45+int(30*index/total),
                        f"Comparaison {index}/{total} : régime {source_regime}")

                con.execute("""INSERT INTO resultats_analyse_multi
                    SELECT uuid(),?,d.ligne_declaratif_id,NULL,?,?,NULL,NULL,NULL,NULL,
                           'AUCUNE','DECLARE_NON_PAYE',0,0,0,0,'AUCUN_IMPACT',
                           d.matricule_normalise,d.nom_normalise
                    FROM declaratif_standardise d
                    WHERE d.institution_id=? AND d.regime=? AND d.trimestre=? AND d.annee=?
                      AND d.execution_id=? AND NOT EXISTS(SELECT 1 FROM resultats_analyse_multi r
                          WHERE r.campagne_id=? AND r.ligne_declaratif_id=d.ligne_declaratif_id)""",
                    [campaign_id,institution_id,declaration_regime,institution_id,
                     declaration_regime,quarter,int(year),declaration_execution_id,campaign_id])
                progress and progress(80,"Classification par institution et régime")
                con.execute("""UPDATE resultats_analyse_multi r SET
                    nombre_occurrences=a.occurrences,nombre_regimes=a.regimes
                    FROM (SELECT ligne_declaratif_id,
                                 COUNT(ligne_paie_id) occurrences,
                                 COUNT(DISTINCT regime_paiement) regimes
                          FROM resultats_analyse_multi WHERE campagne_id=?
                          GROUP BY ligne_declaratif_id) a
                    WHERE r.campagne_id=? AND r.ligne_declaratif_id=a.ligne_declaratif_id""",
                    [campaign_id,campaign_id])
                con.execute("""UPDATE resultats_analyse_multi SET statut_analyse=CASE
                    WHEN ligne_paie_id IS NULL THEN 'DECLARE_NON_PAYE'
                    WHEN nombre_regimes>1 THEN 'PAYE_PLUSIEURS_REGIMES'
                    WHEN regime_paiement<>regime_declaratif
                         AND institution_paiement_id<>institution_declarative_id
                         THEN 'PAYE_AUTRE_INSTITUTION_ET_REGIME'
                    WHEN regime_paiement<>regime_declaratif THEN 'PAYE_AUTRE_REGIME'
                    WHEN institution_paiement_id<>institution_declarative_id
                         THEN 'PAYE_AUTRE_INSTITUTION'
                    WHEN nombre_occurrences>1 THEN 'PAIEMENTS_MULTIPLES_DANS_PERIMETRE'
                    ELSE 'PAYE_DANS_PERIMETRE' END WHERE campagne_id=?""",[campaign_id])
                base_count=con.execute("SELECT COUNT(*) FROM base_analyse_multi WHERE campagne_id=?",
                                       [campaign_id]).fetchone()[0]
                con.execute("""UPDATE campagnes_analyse_multi SET statut='TERMINE',
                    lignes_base=?,termine_le=CURRENT_TIMESTAMP WHERE campagne_id=?""",
                    [base_count,campaign_id])
                con.execute("COMMIT")
            except Exception:
                con.execute("ROLLBACK")
                con.execute("""UPDATE campagnes_analyse_multi SET statut='ERREUR',
                    termine_le=CURRENT_TIMESTAMP WHERE campagne_id=?""",[campaign_id])
                raise
        progress and progress(100,"Analyse multi-régimes terminée")
        return {"campaign_id":campaign_id,"base_rows":base_count,
                "declaration_rows":declaration_count,
                "declaration_execution_id":declaration_execution_id,
                "summary":self.summary(campaign_id)}

    def summary(self,campaign_id: str) -> list[tuple]:
        with self.db.connect() as con:
            return con.execute("""SELECT statut_analyse,COUNT(*),
                    COUNT(DISTINCT COALESCE(NULLIF(matricule_normalise,''),
                        NULLIF(nom_normalise,''),ligne_declaratif_id)),
                    COALESCE(SUM(masse_financiere),0),COALESCE(SUM(impact_potentiel),0)
                FROM resultats_analyse_multi WHERE campagne_id=?
                GROUP BY statut_analyse ORDER BY statut_analyse""",[campaign_id]).fetchall()

    def campaign_details(self,campaign_id: str) -> tuple:
        with self.db.connect() as con:
            row=con.execute("""SELECT c.institution_declarative_id,i.nom_officiel,
                    c.regime_declaratif,c.trimestre,c.annee,c.lignes_base,c.lignes_declaratives
                FROM campagnes_analyse_multi c LEFT JOIN institutions i
                  ON i.institution_id=c.institution_declarative_id
                WHERE c.campagne_id=?""",[campaign_id]).fetchone()
        if not row:raise ValueError("Campagne d'analyse introuvable.")
        return row

    def list_campaigns(self,include_archived: bool=False) -> list[tuple]:
        condition="" if include_archived else "WHERE NOT COALESCE(c.archivee,FALSE)"
        with self.db.connect() as con:
            return con.execute(f"""SELECT c.campagne_id,i.nom_officiel,c.regime_declaratif,
                    c.trimestre,c.annee,c.statut,c.lignes_base,c.lignes_declaratives,
                    c.declaratif_execution_id,c.cree_le,c.termine_le,c.dossier_export,
                    COALESCE(c.archivee,FALSE)
                FROM campagnes_analyse_multi c LEFT JOIN institutions i
                  ON i.institution_id=c.institution_declarative_id
                {condition} ORDER BY c.cree_le DESC""").fetchall()

    def archive_campaign(self,campaign_id: str,archived: bool=True) -> None:
        with self.db.connect() as con:
            found=con.execute("SELECT 1 FROM campagnes_analyse_multi WHERE campagne_id=?",
                              [campaign_id]).fetchone()
            if not found:raise ValueError("Campagne introuvable.")
            con.execute("UPDATE campagnes_analyse_multi SET archivee=? WHERE campagne_id=?",
                        [archived,campaign_id])

    def export(self,campaign_id: str,root: str,progress: Progress=None) -> Path:
        institution_id,institution,regime,quarter,year,base_rows,declaration_rows=self.campaign_details(campaign_id)
        slug=re.sub(r"[^A-Za-z0-9]+","_",institution or institution_id).strip("_").lower()
        folder=Path(root)/f"analyse_multi_{year}_{quarter}_{slug}_{datetime.now():%Y%m%d_%H%M%S}"
        folder.mkdir(parents=True,exist_ok=True)
        progress and progress(10,"Création du rapport synthétique")
        report=Workbook();ws=report.active;ws.title="Synthèse multi-régimes"
        ws.append(["ANALYSE MULTI-RÉGIMES"]);ws["A1"].font=Font(bold=True,size=16,color="12355B")
        ws.append(sanitize_excel_row(["Institution déclarative",institution]));ws.append(sanitize_excel_row(["Périmètre",f"{regime} — {quarter} {year}"]))
        ws.append(["Lignes déclaratives",declaration_rows]);ws.append(["Lignes de la base",base_rows]);ws.append([])
        ws.append(["Catégorie","Enregistrements","Concernés uniques","Masse financière","Impact potentiel"])
        for cell in ws[7]:cell.font=Font(bold=True,color="FFFFFF");cell.fill=PatternFill("solid",fgColor="12355B")
        for row in self.summary(campaign_id):ws.append(sanitize_excel_row(row))
        ws.freeze_panes="A8";ws.auto_filter.ref=f"A7:E{ws.max_row}"
        for col,width in zip("ABCDE",[43,18,20,22,22]):ws.column_dimensions[col].width=width
        methods=report.create_sheet("Méthodologie")
        for line in [
            "Les sources appartiennent strictement au même trimestre et à la même année.",
            "Les filtres existants sont appliqués séparément à chaque institution et régime.",
            "Le matricule est prioritaire; le nom est utilisé seulement sans correspondance matricule.",
            "NU, N.U et leurs variantes sont des matricules non exploitables.",
            "La formule d'impact dépend du régime et de l'institution de paiement.",
            "Les raw_* restent inchangés et chaque ligne conserve son exécution d'origine.",
        ]:methods.append([line])
        methods.column_dimensions["A"].width=115
        report_path=folder/"rapport_multi_regimes.xlsx"

        annex=Workbook(write_only=True);detail=annex.create_sheet("Résultats détaillés")
        detail.append(["Statut","Méthode","Institution déclarative","Régime déclaratif",
            "Institution paiement","Régime paiement","Table source","Matricule","Nom complet","Nom normalisé",
            "Occurrences","Régimes","Masse financière","Impact potentiel","Formule d'impact"])
        with self.db.connect() as con:
            cursor=con.execute("""SELECT r.statut_analyse,r.methode_correspondance,
                    di.nom_officiel,r.regime_declaratif,pi.nom_officiel,r.regime_paiement,
                    r.table_source,r.matricule_normalise,
                    COALESCE(TRIM(CONCAT_WS(' ', COALESCE(NULLIF(b.nom,''),''), COALESCE(NULLIF(b.prenom,''),''))), r.nom_normalise),
                    r.nom_normalise,r.nombre_occurrences,
                    r.nombre_regimes,r.masse_financiere,r.impact_potentiel,r.formule_impact_id
                FROM resultats_analyse_multi r
                LEFT JOIN institutions di ON di.institution_id=r.institution_declarative_id
                LEFT JOIN institutions pi ON pi.institution_id=r.institution_paiement_id
                LEFT JOIN base_analyse_multi b
                  ON b.campagne_id=r.campagne_id AND b.ligne_paie_id=r.ligne_paie_id
                WHERE r.campagne_id=?
                ORDER BY r.statut_analyse,r.regime_paiement,r.nom_normalise""",[campaign_id])
            count=0
            while True:
                rows=cursor.fetchmany(5000)
                if not rows:break
                for row in rows:detail.append(sanitize_excel_row(row))
                count+=len(rows)
                progress and progress(min(90,35+int(55*count/max(1,base_rows+declaration_rows))),
                    f"Annexe détaillée : {count:,} lignes écrites".replace(","," "))
            sources=annex.create_sheet("Sources et filtres")
            sources.append(["Exécution","Institution","Régime","Table","Disponibles","Retenues","Filtres"])
            for row in con.execute("""SELECT s.execution_id,i.nom_officiel,s.regime,s.table_source,
                    s.lignes_disponibles,s.lignes_retenues,s.filtres_appliques
                FROM sources_analyse_multi s LEFT JOIN institutions i
                  ON i.institution_id=s.institution_id
                WHERE s.campagne_id=? ORDER BY s.regime,i.nom_officiel""",[campaign_id]).fetchall():
                sources.append(sanitize_excel_row(row))
        annex_path=folder/"annexe_resultats_multi_regimes.xlsx";annex.save(annex_path)
        progress and progress(72,f"Fichier généré : {annex_path.name}")

        category_folder=folder/"annexes_par_categorie";category_folder.mkdir(exist_ok=True)
        categories=[row[0] for row in self.summary(campaign_id)]
        category_headers=["Méthode","Institution déclarative","Régime déclaratif",
            "Institution paiement","Régime paiement","Table source","Matricule","Nom complet","Nom normalisé",
            "Occurrences","Nombre de régimes","Masse financière","Impact potentiel","Formule"]
        for index,status in enumerate(categories,1):
            book=Workbook(write_only=True);sheet=book.create_sheet("Détails");sheet.append(category_headers)
            with self.db.connect() as con:
                cursor=con.execute("""SELECT r.methode_correspondance,di.nom_officiel,
                        r.regime_declaratif,pi.nom_officiel,r.regime_paiement,r.table_source,
                        r.matricule_normalise,
                        COALESCE(TRIM(CONCAT_WS(' ', COALESCE(NULLIF(b.nom,''),''), COALESCE(NULLIF(b.prenom,''),''))), r.nom_normalise),
                        r.nom_normalise,r.nombre_occurrences,r.nombre_regimes,
                        r.masse_financiere,r.impact_potentiel,r.formule_impact_id
                    FROM resultats_analyse_multi r
                    LEFT JOIN institutions di ON di.institution_id=r.institution_declarative_id
                    LEFT JOIN institutions pi ON pi.institution_id=r.institution_paiement_id
                    LEFT JOIN base_analyse_multi b
                      ON b.campagne_id=r.campagne_id AND b.ligne_paie_id=r.ligne_paie_id
                    WHERE r.campagne_id=? AND r.statut_analyse=?
                    ORDER BY r.regime_paiement,r.nom_normalise""",[campaign_id,status])
                while True:
                    rows=cursor.fetchmany(5000)
                    if not rows:break
                    for row in rows:sheet.append(sanitize_excel_row(row))
            filename=re.sub(r"[^A-Za-z0-9_]+","_",status).lower()+".xlsx"
            book.save(category_folder/filename)
            progress and progress(72+int(10*index/max(1,len(categories))),
                f"Fichier généré : annexes_par_categorie/{filename}")

        effectif=Workbook(write_only=True);unique=effectif.create_sheet("Effectifs uniques")
        unique.append(["Matricule","Nom complet","Nom normalisé","Statut","Régimes de paiement",
                       "Institutions de paiement","Occurrences","Masse financière","Impact potentiel"])
        with self.db.connect() as con:
            cursor=con.execute("""SELECT r.matricule_normalise,
                    COALESCE(TRIM(CONCAT_WS(' ', COALESCE(NULLIF(b.nom,''),''), COALESCE(NULLIF(b.prenom,''),''))), r.nom_normalise),
                    r.nom_normalise,MAX(statut_analyse),STRING_AGG(DISTINCT COALESCE(regime_paiement,'NON PAYE'),', '),
                    STRING_AGG(DISTINCT COALESCE(pi.nom_officiel,'NON PAYE'),', '),
                    MAX(nombre_occurrences),SUM(masse_financiere),SUM(impact_potentiel)
                FROM resultats_analyse_multi r
                LEFT JOIN institutions pi ON pi.institution_id=r.institution_paiement_id
                LEFT JOIN base_analyse_multi b
                  ON b.campagne_id=r.campagne_id AND b.ligne_paie_id=r.ligne_paie_id
                WHERE r.campagne_id=? GROUP BY r.ligne_declaratif_id,
                    r.matricule_normalise,r.nom_normalise,COALESCE(TRIM(CONCAT_WS(' ', COALESCE(NULLIF(b.nom,''),''), COALESCE(NULLIF(b.prenom,''),''))), r.nom_normalise)
                ORDER BY r.nom_normalise""",[campaign_id])
            while True:
                rows=cursor.fetchmany(5000)
                if not rows:break
                for row in rows:unique.append(sanitize_excel_row(row))
        effectif_path=folder/"effectifs_uniques_multi_regimes.xlsx";effectif.save(effectif_path)
        progress and progress(90,f"Fichier généré : {effectif_path.name}")

        letter=Document();normal=letter.styles["Normal"];normal.font.name="Arial";normal.font.size=Pt(10)
        title=letter.add_paragraph();run=title.add_run("LETTRE D’INTERPRÉTATION — ANALYSE MULTI-RÉGIMES")
        run.bold=True;run.font.size=Pt(15)
        letter.add_paragraph(sanitize_xml_text(f"Institution concernée : {institution}"))
        letter.add_paragraph(sanitize_xml_text(f"Période analysée : {quarter} {year} — régime déclaratif {regime}"))
        letter.add_paragraph(
            f"La campagne a comparé {declaration_rows:,} lignes déclaratives à une base constituée "
            f"de {base_rows:,} lignes de paie, après application séparée des filtres métier "
            "de chaque institution et régime.".replace(","," "))
        table=letter.add_table(rows=1,cols=4);table.style="Table Grid"
        for cell,value in zip(table.rows[0].cells,["Catégorie","Concernés","Masse","Impact potentiel"]):
            cell.text=value
        for status,_records,people,mass,impact in self.summary(campaign_id):
            cells=table.add_row().cells
            for cell,value in zip(cells,[status,people,f"{mass:,.2f}".replace(","," "),
                                         f"{impact:,.2f}".replace(","," ")]):cell.text=sanitize_xml_text(value)
        letter.add_paragraph(
            "Les impacts sont calculés selon le régime et l’institution de la ligne de paiement "
            "retrouvée. Ils constituent des impacts potentiels à confirmer par les pièces "
            "administratives et la validation métier.")
        letter.add_paragraph(
            "Les fichiers joints détaillent les sources, les filtres appliqués, les correspondances "
            "par matricule ou par nom et les effectifs uniques concernés.")
        letter_path=folder/"lettre_interpretation_multi_regimes.docx";letter.save(letter_path)
        progress and progress(97,f"Fichier généré : {letter_path.name}")
        ws.append([]);ws.append(["Liens annexes","Annexe détaillée","Effectifs uniques","Lettre d’interprétation"])
        links_row=ws.max_row
        for col,target in [(2,"annexe_resultats_multi_regimes.xlsx"),(3,"effectifs_uniques_multi_regimes.xlsx"),(4,"lettre_interpretation_multi_regimes.docx")]:
            cell=ws.cell(links_row,col);cell.value=target;cell.hyperlink=target;cell.style="Hyperlink"
        for cell in ws[links_row]:
            cell.font=Font(bold=True,color="12355B")
        report.save(report_path)
        progress and progress(35,f"Fichier généré : {report_path.name}")
        with self.db.connect() as con:
            con.execute("UPDATE campagnes_analyse_multi SET dossier_export=? WHERE campagne_id=?",
                        [str(folder),campaign_id])
        progress and progress(100,"Rapport, annexes par catégorie, effectifs et lettre terminés")
        return folder
